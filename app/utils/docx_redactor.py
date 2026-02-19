from docx import Document
from io import BytesIO

def redact_docx_paragraphwise(
    original_doc_bytes: bytes,
    pipeline,
    selected_entities: list[str] | None
) -> tuple[BytesIO, int]:
    doc = Document(BytesIO(original_doc_bytes))
    total_entity_count = 0

    for para in doc.paragraphs:
        full_text = "".join([run.text for run in para.runs])

        if not full_text:
            continue

        _, entities = pipeline.run(full_text)
        
        chars_to_redact = [False] * len(full_text)
        
        for e in entities:
            if selected_entities is not None and e.entity_type not in selected_entities:
                continue

            if e.start < 0 or e.end > len(full_text):
                 continue
                 
            span_len = e.end - e.start
            if span_len > 60 or (len(full_text) > 50 and span_len / len(full_text) > 0.8):
                 continue

            total_entity_count += 1

            for i in range(e.start, e.end):
                chars_to_redact[i] = True

        current_global_idx = 0
        for run in para.runs:
            run_chars = list(run.text)
            for i in range(len(run_chars)):
                if current_global_idx < len(chars_to_redact) and chars_to_redact[current_global_idx]:
                    run_chars[i] = "*"
                current_global_idx += 1
            run.text = "".join(run_chars)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, total_entity_count

def redact_docx_preview(
    original_doc_bytes: bytes,
    pipeline,
    selected_entities: list[str] | None,
    limit_paragraphs: int = 20
) -> str:
    doc = Document(BytesIO(original_doc_bytes))
    preview_text = ""
    
    count = 0
    for para in doc.paragraphs:
        if count >= limit_paragraphs:
            break
            
        full_text = "".join([run.text for run in para.runs])
        if not full_text.strip():
            continue
            
        _, entities = pipeline.run(full_text)
        
        chars = list(full_text)
        
        for e in entities:
             if selected_entities is not None and e.entity_type not in selected_entities:
                 continue

             if e.start < 0 or e.end > len(full_text):
                  continue
                  
             span_len = e.end - e.start
             if span_len > 60 or (len(full_text) > 50 and span_len / len(full_text) > 0.8):
                  continue

             for i in range(e.start, e.end):
                 chars[i] = "*"
        
        preview_text += "".join(chars) + "\n\n"
        count += 1
        
    return preview_text
